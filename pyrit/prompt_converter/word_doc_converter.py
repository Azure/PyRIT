# Copyright (c) Microsoft Corporation.
# Licensed under the MIT license.

from __future__ import annotations

import ast
import hashlib
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document

from pyrit.common.logger import logger
from pyrit.identifiers import ComponentIdentifier
from pyrit.models import PromptDataType, SeedPrompt, data_serializer_factory
from pyrit.models.data_type_serializer import DataTypeSerializer
from pyrit.prompt_converter.prompt_converter import ConverterResult, PromptConverter


@dataclass
class _WordDocInjectionConfig:
    """Configuration for how to inject content into a Word document."""

    existing_docx: Optional[Path]
    placeholder: str


class WordDocConverter(PromptConverter):
    """
    Convert a text prompt into a Word (.docx) document.

    This converter supports two main modes:

    1. **New document generation**
       If no existing document is provided, the converter creates a simple `.docx`
       containing the rendered prompt content in a single paragraph.

    2. **Placeholder-based injection into an existing document**
       If an ``existing_docx`` is provided, the converter searches for a literal
       placeholder string (for example ``{{INJECTION_PLACEHOLDER}}``) in the
       document's paragraphs. When the placeholder is found fully inside a single
       run, it is replaced with the rendered prompt content while preserving the
       rest of the paragraph and its formatting.

       .. important::
          Placeholders must be fully contained within a single run. If a
          placeholder spans multiple runs (for example due to mixed formatting),
          this converter will not replace it. This limitation is intentional to
          avoid collapsing mixed formatting or rewriting complex run structures.

    Security note:
        This converter does **not** render Jinja2 templates from arbitrary
        ``.docx`` content. Templating is handled via ``SeedPrompt`` (if provided),
        and only the already-rendered text is injected into the document. This
        avoids executing untrusted Jinja2 templates from document bodies.
    """

    SUPPORTED_INPUT_TYPES = ("text",)
    SUPPORTED_OUTPUT_TYPES = ("binary_path",)

    def __init__(
        self,
        *,
        prompt_template: Optional[SeedPrompt] = None,
        existing_docx: Optional[Path] = None,
        placeholder: str = "{{INJECTION_PLACEHOLDER}}",
    ) -> None:
        """
        Initialize the Word document converter.

        Args:
            prompt_template: Optional ``SeedPrompt`` template used to render the
                final content before injection. If provided, ``prompt`` passed
                to ``convert_async`` must be a string whose contents can be
                interpreted as the template parameters (for example, a
                JSON-encoded or other parseable mapping of keys to values).
            existing_docx: Optional path to an existing `.docx` file. When
                provided, the converter will search for ``placeholder`` inside
                the document paragraphs and replace it with the rendered content.
                If not provided, a new document is generated instead.
            placeholder: Literal placeholder text to search for in the existing
                document. This value must be fully contained within a single
                run for the replacement to succeed.

        Raises:
            FileNotFoundError: If ``existing_docx`` is provided but does not exist.
            ValueError: If ``placeholder`` is empty.
        """
        super().__init__()

        if not placeholder:
            raise ValueError("Placeholder must be a non-empty string.")

        if existing_docx is not None and not existing_docx.is_file():
            raise FileNotFoundError(f"Word document not found at: {existing_docx}")

        self._prompt_template = prompt_template
        self._injection_config = _WordDocInjectionConfig(
            existing_docx=existing_docx,
            placeholder=placeholder,
        )

    def _build_identifier(self) -> ComponentIdentifier:
        """
        Build identifier with template and document parameters.

        Returns:
            ComponentIdentifier: The identifier with converter-specific parameters.
        """
        template_hash: Optional[str] = None
        if self._prompt_template:
            template_hash = hashlib.sha256(str(self._prompt_template.value).encode("utf-8")).hexdigest()[:16]

        existing_docx_path = None
        if self._injection_config.existing_docx:
            existing_docx_path = str(self._injection_config.existing_docx)

        return self._create_identifier(
            params={
                "prompt_template_hash": template_hash,
                "existing_docx_path": existing_docx_path,
                "placeholder": self._injection_config.placeholder,
            }
        )

    async def convert_async(self, *, prompt: str, input_type: PromptDataType = "text") -> ConverterResult:
        """
        Convert the given prompt into a Word document (.docx).

        If ``prompt_template`` is provided, the prompt is first used to render the
        template via ``SeedPrompt.render_template_value``. Otherwise, the raw
        ``prompt`` string is used as the content.

        - When ``existing_docx`` is set, this content is injected into the
          document by replacing the configured placeholder string.
        - When no ``existing_docx`` is provided, a new document with a single
          paragraph containing the content is created.

        Args:
            prompt: The prompt or dynamic data used to generate the content.
            input_type: The type of input data. Must be ``"text"``.

        Returns:
            ConverterResult: Contains the path to the generated `.docx` file in
            ``output_text`` and ``output_type="binary_path"``.

        Raises:
            ValueError: If the input type is not supported.
        """
        if not self.input_supported(input_type):
            raise ValueError("Input type not supported")

        content = self._prepare_content(prompt)

        if self._injection_config.existing_docx:
            doc_bytes = self._inject_into_existing_docx(content)
        else:
            doc_bytes = self._generate_new_docx(content)

        serializer = await self._serialize_docx_async(doc_bytes)

        return ConverterResult(output_text=serializer.value, output_type="binary_path")

    def _prepare_content(self, prompt: str) -> str:
        """
        Prepare the content to be injected or written to the document.

        If a ``SeedPrompt`` template is provided, the ``prompt`` is parsed (if
        necessary) as a dictionary and used to render the template. Otherwise,
        the raw prompt string is used.

        Returns:
            str: The rendered or raw content string.

        Raises:
            ValueError: If the prompt cannot be parsed or rendered with the template.
        """
        if self._prompt_template:
            logger.debug(f"Preparing Word content with template: {self._prompt_template.value}")
            try:
                dynamic_data: Dict[str, Any] = ast.literal_eval(prompt)

                if not isinstance(dynamic_data, dict):
                    raise ValueError("Prompt must be a dictionary-compatible object after parsing.")

                rendered_content = self._prompt_template.render_template_value(**dynamic_data)
                logger.debug("Rendered Word template content successfully.")
                return rendered_content
            except (ValueError, SyntaxError, KeyError) as exc:
                logger.error("Error rendering Word template content: %s", exc)
                raise ValueError(f"Failed to render the prompt for Word document: {exc}") from exc

        logger.debug("No template provided for Word document. Using raw prompt content.")
        return prompt

    def _generate_new_docx(self, content: str) -> bytes:
        """
        Generate a new `.docx` document containing the given content.

        Returns:
            bytes: The serialized `.docx` document bytes.
        """
        document = Document()
        document.add_paragraph(content)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _inject_into_existing_docx(self, content: str) -> bytes:
        """
        Inject content into an existing document by replacing the placeholder.

        The placeholder must appear fully inside a single run; if it only exists
        across multiple runs, it will not be replaced.

        Returns:
            bytes: The serialized `.docx` document bytes.

        Raises:
            ValueError: If no existing_docx path was provided.
        """
        if self._injection_config.existing_docx is None:
            raise ValueError(
                "Cannot inject into an existing Word document because no 'existing_docx' "
                "path was provided in the injection configuration."
            )
        document = Document(str(self._injection_config.existing_docx))

        placeholder = self._injection_config.placeholder
        replaced_any = False

        for paragraph in document.paragraphs:
            if placeholder not in paragraph.text:
                continue

            if self._replace_placeholder_in_paragraph(paragraph, placeholder, content):
                replaced_any = True

        if not replaced_any:
            logger.warning(
                "No placeholder '%s' found in document '%s' or placeholder spanned multiple runs.",
                placeholder,
                self._injection_config.existing_docx,
            )

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def _replace_placeholder_in_paragraph(paragraph: Any, placeholder: str, content: str) -> bool:
        """
        Replace a placeholder inside a single run of a paragraph.

        This function searches all runs of a paragraph and performs a string
        replacement in the first run whose text contains the placeholder. It
        does not modify other runs, which helps preserve existing formatting.

        Returns:
            bool: True if a replacement was made, False otherwise.
        """
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, content)
                return True
        return False

    async def _serialize_docx_async(self, docx_bytes: bytes) -> DataTypeSerializer:
        """
        Serialize the generated document using a data serializer.

        Returns:
            DataTypeSerializer: The serializer containing the saved document path.
        """
        extension = "docx"

        serializer = data_serializer_factory(
            category="prompt-memory-entries",
            data_type="binary_path",
            extension=extension,
        )
        await serializer.save_data(docx_bytes)
        return serializer
