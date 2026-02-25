# Copyright (c) Microsoft Corporation.
# Licensed under the MIT license.

from io import BytesIO
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from pyrit.models import SeedPrompt
from pyrit.prompt_converter import ConverterResult, WordDocConverter


@pytest.fixture
def word_doc_converter_no_template() -> WordDocConverter:
    """A WordDocConverter with no template or existing document."""
    return WordDocConverter(prompt_template=None)


def test_input_supported(word_doc_converter_no_template: WordDocConverter) -> None:
    """Test that only 'text' input types are supported."""
    assert word_doc_converter_no_template.input_supported("text") is True
    assert word_doc_converter_no_template.input_supported("image") is False


@pytest.mark.asyncio
async def test_convert_async_creates_new_docx(word_doc_converter_no_template: WordDocConverter) -> None:
    """convert_async should generate a new document when no existing_docx is configured."""
    prompt = "Hello, Word!"
    mock_doc_bytes = b"mock-docx-bytes"

    with (
        patch.object(word_doc_converter_no_template, "_prepare_content", return_value=prompt) as mock_prepare,
        patch.object(
            word_doc_converter_no_template,
            "_generate_new_docx",
            return_value=mock_doc_bytes,
        ) as mock_generate,
        patch.object(word_doc_converter_no_template, "_inject_into_existing_docx") as mock_inject,
        patch.object(word_doc_converter_no_template, "_serialize_docx_async") as mock_serialize,
    ):
        serializer_mock = MagicMock()
        serializer_mock.value = "mock_path.docx"
        mock_serialize.return_value = serializer_mock

        result = await word_doc_converter_no_template.convert_async(prompt=prompt)

    mock_prepare.assert_called_once_with(prompt)
    mock_generate.assert_called_once_with(prompt)
    mock_inject.assert_not_called()
    mock_serialize.assert_called_once_with(mock_doc_bytes)

    assert isinstance(result, ConverterResult)
    assert result.output_type == "binary_path"
    assert result.output_text == "mock_path.docx"


@pytest.mark.asyncio
async def test_convert_async_injects_into_existing_docx(tmp_path: Path) -> None:
    """convert_async should inject into an existing document when existing_docx is configured."""
    existing_docx_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Prefix {{INJECTION_PLACEHOLDER}} Suffix")
    document.save(existing_docx_path)

    converter = WordDocConverter(existing_docx=existing_docx_path)

    prompt = "Injected content"
    mock_doc_bytes = b"modified-docx-bytes"

    with (
        patch.object(converter, "_prepare_content", return_value=prompt) as mock_prepare,
        patch.object(converter, "_generate_new_docx") as mock_generate,
        patch.object(converter, "_inject_into_existing_docx", return_value=mock_doc_bytes) as mock_inject,
        patch.object(converter, "_serialize_docx_async") as mock_serialize,
    ):
        serializer_mock = MagicMock()
        serializer_mock.value = "injected.docx"
        mock_serialize.return_value = serializer_mock

        result = await converter.convert_async(prompt=prompt)

    mock_prepare.assert_called_once_with(prompt)
    mock_generate.assert_not_called()
    mock_inject.assert_called_once_with(prompt)
    mock_serialize.assert_called_once_with(mock_doc_bytes)

    assert isinstance(result, ConverterResult)
    assert result.output_type == "binary_path"
    assert result.output_text == "injected.docx"


@pytest.mark.asyncio
async def test_convert_async_unsupported_input_type(word_doc_converter_no_template: WordDocConverter) -> None:
    """convert_async should raise for unsupported input types."""
    with pytest.raises(ValueError, match="Input type not supported"):
        await word_doc_converter_no_template.convert_async(prompt="ignored", input_type="image")


def test_inject_into_existing_docx_requires_existing_path(tmp_path: Path) -> None:
    """_inject_into_existing_docx should raise if called without an existing_docx path."""
    converter = WordDocConverter(existing_docx=None)

    with pytest.raises(ValueError, match="no 'existing_docx' path was provided"):
        converter._inject_into_existing_docx("content")


def test_placeholder_replaced_when_in_single_run(tmp_path: Path) -> None:
    """The placeholder should be replaced when it appears in a single run."""
    existing_docx_path = tmp_path / "placeholder_single_run.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hello {{INJECTION_PLACEHOLDER}}!")
    document.save(existing_docx_path)

    converter = WordDocConverter(existing_docx=existing_docx_path)

    result_bytes = converter._inject_into_existing_docx("WORLD")
    new_doc = Document(BytesIO(result_bytes))

    assert "WORLD" in new_doc.paragraphs[0].text
    assert "{{INJECTION_PLACEHOLDER}}" not in new_doc.paragraphs[0].text


def test_placeholder_not_replaced_when_split_across_runs(tmp_path: Path) -> None:
    """The placeholder should not be replaced when split across multiple runs."""
    existing_docx_path = tmp_path / "placeholder_multi_run.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hello {{")
    paragraph.add_run("INJECTION_PLACEHOLDER}}!")
    document.save(existing_docx_path)

    converter = WordDocConverter(existing_docx=existing_docx_path)

    result_bytes = converter._inject_into_existing_docx("WORLD")
    new_doc = Document(BytesIO(result_bytes))

    # Since the placeholder spans multiple runs, the content should be unchanged.
    assert new_doc.paragraphs[0].text == "Hello {{INJECTION_PLACEHOLDER}}!"


def test_constructor_raises_on_empty_placeholder() -> None:
    """Constructor should raise ValueError when placeholder is empty."""
    with pytest.raises(ValueError, match="Placeholder must be a non-empty string"):
        WordDocConverter(placeholder="")


def test_constructor_raises_on_nonexistent_file(tmp_path: Path) -> None:
    """Constructor should raise FileNotFoundError when existing_docx doesn't exist."""
    fake_path = tmp_path / "nonexistent.docx"
    with pytest.raises(FileNotFoundError, match="Word document not found at"):
        WordDocConverter(existing_docx=fake_path)


def test_constructor_accepts_valid_existing_docx(tmp_path: Path) -> None:
    """Constructor should accept a valid existing docx file."""
    docx_path = tmp_path / "valid.docx"
    Document().save(docx_path)
    converter = WordDocConverter(existing_docx=docx_path)
    assert converter._injection_config.existing_docx == docx_path


def test_constructor_custom_placeholder() -> None:
    """Constructor should accept a custom placeholder string."""
    converter = WordDocConverter(placeholder="<<REPLACE_ME>>")
    assert converter._injection_config.placeholder == "<<REPLACE_ME>>"


def test_build_identifier_without_template() -> None:
    """_build_identifier should return correct params when no template is set."""
    converter = WordDocConverter()
    identifier = converter._build_identifier()
    assert "prompt_template_hash" not in identifier.params
    assert "existing_docx_path" not in identifier.params
    assert identifier.params["placeholder"] == "{{INJECTION_PLACEHOLDER}}"


def test_build_identifier_with_template() -> None:
    """_build_identifier should include a template hash when template is set."""
    template = SeedPrompt(value="Hello {{ name }}", data_type="text", name="test", parameters=["name"])
    converter = WordDocConverter(prompt_template=template)
    identifier = converter._build_identifier()
    assert identifier.params["prompt_template_hash"] is not None
    assert len(identifier.params["prompt_template_hash"]) == 16


def test_build_identifier_with_existing_docx(tmp_path: Path) -> None:
    """_build_identifier should include the docx path when set."""
    docx_path = tmp_path / "test.docx"
    Document().save(docx_path)
    converter = WordDocConverter(existing_docx=docx_path)
    identifier = converter._build_identifier()
    assert identifier.params["existing_docx_path"] == str(docx_path)


def test_prepare_content_raw_string() -> None:
    """_prepare_content should return the raw prompt when no template is set."""
    converter = WordDocConverter()
    result = converter._prepare_content("Hello world")
    assert result == "Hello world"


def test_prepare_content_with_template() -> None:
    """_prepare_content should render template with prompt data."""
    template = SeedPrompt(value="Dear {{ name }}, welcome!", data_type="text", name="test", parameters=["name"])
    converter = WordDocConverter(prompt_template=template)
    result = converter._prepare_content("{'name': 'Alice'}")
    assert result == "Dear Alice, welcome!"


def test_prepare_content_template_invalid_prompt() -> None:
    """_prepare_content should raise ValueError when prompt can't be parsed as dict."""
    template = SeedPrompt(value="Hello {{ name }}", data_type="text", name="test", parameters=["name"])
    converter = WordDocConverter(prompt_template=template)
    with pytest.raises(ValueError, match="Failed to render the prompt for Word document"):
        converter._prepare_content("not a dict")


def test_prepare_content_template_missing_key() -> None:
    """_prepare_content should raise ValueError when template key is missing from prompt."""
    template = SeedPrompt(value="Hello {{ name }}", data_type="text", name="test", parameters=["name"])
    converter = WordDocConverter(prompt_template=template)
    with pytest.raises(ValueError, match="Failed to render the prompt for Word document"):
        converter._prepare_content("{'wrong_key': 'value'}")


def test_generate_new_docx_creates_valid_document() -> None:
    """_generate_new_docx should produce a valid docx with the given content."""
    converter = WordDocConverter()
    content = "Test paragraph content"
    result_bytes = converter._generate_new_docx(content)

    doc = Document(BytesIO(result_bytes))
    assert len(doc.paragraphs) > 0
    assert doc.paragraphs[0].text == content


def test_inject_with_custom_placeholder(tmp_path: Path) -> None:
    """Injection should work with a custom placeholder string."""
    existing_docx_path = tmp_path / "custom_placeholder.docx"
    document = Document()
    document.add_paragraph("Before <<REPLACE>> after")
    document.save(existing_docx_path)

    converter = WordDocConverter(existing_docx=existing_docx_path, placeholder="<<REPLACE>>")
    result_bytes = converter._inject_into_existing_docx("INJECTED")
    new_doc = Document(BytesIO(result_bytes))

    assert new_doc.paragraphs[0].text == "Before INJECTED after"


def test_inject_multiple_placeholders_in_document(tmp_path: Path) -> None:
    """All placeholders across paragraphs should be replaced."""
    existing_docx_path = tmp_path / "multi_placeholder.docx"
    document = Document()
    document.add_paragraph("First {{INJECTION_PLACEHOLDER}} here")
    document.add_paragraph("Second {{INJECTION_PLACEHOLDER}} here")
    document.save(existing_docx_path)

    converter = WordDocConverter(existing_docx=existing_docx_path)
    result_bytes = converter._inject_into_existing_docx("REPLACED")
    new_doc = Document(BytesIO(result_bytes))

    assert "REPLACED" in new_doc.paragraphs[0].text
    assert "REPLACED" in new_doc.paragraphs[1].text
    assert "{{INJECTION_PLACEHOLDER}}" not in new_doc.paragraphs[0].text
    assert "{{INJECTION_PLACEHOLDER}}" not in new_doc.paragraphs[1].text


def test_inject_no_placeholder_in_document_logs_warning(tmp_path: Path) -> None:
    """When no placeholder is found, a warning should be logged."""
    existing_docx_path = tmp_path / "no_placeholder.docx"
    document = Document()
    document.add_paragraph("No placeholder here")
    document.save(existing_docx_path)

    converter = WordDocConverter(existing_docx=existing_docx_path)

    with patch("pyrit.prompt_converter.word_doc_converter.logger") as mock_logger:
        converter._inject_into_existing_docx("content")
        mock_logger.warning.assert_called_once()


@pytest.mark.asyncio
async def test_convert_async_end_to_end_new_doc(sqlite_instance) -> None:
    """End-to-end: convert_async creates a valid docx file (no mocking of internal logic)."""
    converter = WordDocConverter()
    result = await converter.convert_async(prompt="End-to-end test content")

    assert isinstance(result, ConverterResult)
    assert result.output_type == "binary_path"
    assert result.output_text  # non-empty path


@pytest.mark.asyncio
async def test_convert_async_end_to_end_injection(tmp_path: Path, sqlite_instance) -> None:
    """End-to-end: convert_async injects into an existing docx (no mocking of internal logic)."""
    existing_docx_path = tmp_path / "e2e_template.docx"
    document = Document()
    document.add_paragraph("Hello {{INJECTION_PLACEHOLDER}} world")
    document.save(existing_docx_path)

    converter = WordDocConverter(existing_docx=existing_docx_path)
    result = await converter.convert_async(prompt="INJECTED_VALUE")

    assert isinstance(result, ConverterResult)
    assert result.output_type == "binary_path"
    assert result.output_text  # non-empty path


def test_replace_placeholder_in_paragraph_returns_false_when_not_found() -> None:
    """_replace_placeholder_in_paragraph should return False when placeholder is not in any run."""
    document = Document()
    paragraph = document.add_paragraph("No placeholder here")
    result = WordDocConverter._replace_placeholder_in_paragraph(paragraph, "{{INJECTION_PLACEHOLDER}}", "content")
    assert result is False


def test_replace_placeholder_in_paragraph_returns_true_on_match() -> None:
    """_replace_placeholder_in_paragraph should return True and replace content."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Text with {{INJECTION_PLACEHOLDER}} inside")
    result = WordDocConverter._replace_placeholder_in_paragraph(paragraph, "{{INJECTION_PLACEHOLDER}}", "REPLACED")
    assert result is True
    assert paragraph.runs[0].text == "Text with REPLACED inside"
